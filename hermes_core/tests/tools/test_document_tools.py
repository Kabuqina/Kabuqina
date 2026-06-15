from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def _fake_render_callback(captured: dict, *, slide_count: int = 0):
    """Stand in for the desk clarify callback that drives PptxGenJS in the webview."""
    import base64

    def _cb(question, choices, kind=None, artifact=None):
        captured["question"] = question
        captured["kind"] = kind
        captured["artifact"] = artifact
        payload = base64.b64encode(b"PK\x03\x04 fake pptx bytes").decode()
        deck_slides = ((artifact or {}).get("deck") or {}).get("slides") or []
        return {
            "action": "rendered",
            "text": "",
            "data": {"pptx_base64": payload, "slide_count": slide_count or (len(deck_slides) + 1)},
        }

    return _cb


def test_pdf_read_precise_rejects_paths_outside_workspace(tmp_path, monkeypatch):
    from tools.document_tools import pdf_read_precise

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    outside = tmp_path / "outside.pdf"
    outside.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(workspace))
    result = json.loads(pdf_read_precise(path=str(outside)))

    assert result.get("code") == "outside_workspace"
    assert "workspace" in result
    assert "terminal" in result.get("hint", "").lower()


def test_pptx_write_emits_deck_spec_and_writes_bytes(tmp_path):
    from tools.document_tools import pptx_write

    out = tmp_path / "student-report.pptx"
    captured: dict = {}
    result = json.loads(
        pptx_write(
            path=str(out),
            title="课设答辩",
            slides=[
                {"title": "项目目标", "bullets": ["完成 PDF 精确识别", "生成可提交 PPT"]},
                {"title": "实现方案", "bullets": ["Docling 解析", "大纲确认后生成"]},
            ],
            template="code_defense",
            visual_master="neo_grid_bold",
            callback=_fake_render_callback(captured),
        )
    )

    # The tool hands the deck to the webview as a pptx_render interaction.
    assert captured["kind"] == "pptx_render"
    deck = captured["artifact"]["deck"]
    assert deck["title"] == "课设答辩"
    assert deck["template"] == "code_defense"
    assert deck["visual_master"] == "neo_grid_bold"
    assert [s["title"] for s in deck["slides"]] == ["项目目标", "实现方案"]
    assert deck["slides"][0]["bullets"] == ["完成 PDF 精确识别", "生成可提交 PPT"]

    # The returned base64 is decoded and persisted by Python.
    assert result["ok"] is True
    assert result["slide_count"] == 3
    assert result["template"] == "code_defense"
    assert result["theme"] == "项目答辩"
    assert result["visual_master"] == "neo_grid_bold"
    assert result["visual_master_renderer"] == "pptxgenjs_v1"
    assert out.exists()
    assert out.read_bytes().startswith(b"PK")


def test_deck_meta_normalized_onto_deck_spec():
    from tools.document_tools import _build_deck_spec, _get_pptx_theme

    deck = _build_deck_spec(
        "基于大数据的社交媒体分析",
        [{"title": "研究背景", "bullets": ["信息爆炸"]}],
        _get_pptx_theme("paper_report"),
        "soft_editorial",
        meta={
            "author": "刘志红",
            "affiliation": "河北东方学院",
            "date": "2024",
            "citation": "刘志红.基于大数据的社交媒体分析[J].软件,2024,45(03):183-186",
            "junk": "ignored",  # unknown keys are dropped
        },
    )
    assert deck["meta"] == {
        "author": "刘志红",
        "affiliation": "河北东方学院",
        "date": "2024",
        "citation": "刘志红.基于大数据的社交媒体分析[J].软件,2024,45(03):183-186",
    }


def test_deck_meta_defaults_to_empty_when_absent():
    from tools.document_tools import _build_deck_spec, _get_pptx_theme

    deck = _build_deck_spec("课设答辩", [], _get_pptx_theme("code_defense"), "neo_grid_bold")
    assert deck["meta"] == {}


def _write_theme_pptx(path: Path) -> None:
    """Minimal .pptx-shaped zip carrying just the theme part the extractor reads."""
    import zipfile

    theme_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="School">'
        '<a:themeElements><a:clrScheme name="School">'
        '<a:dk1><a:sysClr val="windowText" lastClr="111111"/></a:dk1>'
        '<a:lt1><a:srgbClr val="FDFAE7"/></a:lt1>'
        '<a:dk2><a:srgbClr val="1E2BFA"/></a:dk2>'
        '<a:lt2><a:srgbClr val="EEEEEE"/></a:lt2>'
        '<a:accent1><a:srgbClr val="E6FF3D"/></a:accent1>'
        '<a:accent2><a:srgbClr val="059669"/></a:accent2>'
        "</a:clrScheme>"
        '<a:fontScheme name="School">'
        '<a:majorFont><a:latin typeface="Source Han Sans"/></a:majorFont>'
        '<a:minorFont><a:latin typeface="Calibri"/></a:minorFont>'
        "</a:fontScheme>"
        "<a:fmtScheme/></a:themeElements></a:theme>"
    )
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("ppt/theme/theme1.xml", theme_xml)


def test_extract_pptx_theme_maps_colors_and_fonts(tmp_path):
    from tools.document_tools import _extract_pptx_theme

    tpl = tmp_path / "school.pptx"
    _write_theme_pptx(tpl)
    palette = _extract_pptx_theme(tpl)
    assert palette == {
        "background": "FDFAE7",  # lt1
        "title": "1E2BFA",       # dk2 (brand dark)
        "body": "111111",        # dk1 sysClr lastClr
        "accent": "E6FF3D",      # accent1
        "accent2": "059669",     # accent2
        "fonts": {"major": "Source Han Sans", "minor": "Calibri"},
    }


def test_extract_pptx_theme_returns_none_on_bad_file(tmp_path):
    from tools.document_tools import _extract_pptx_theme

    bad = tmp_path / "not-a-deck.pptx"
    bad.write_bytes(b"this is not a zip")
    assert _extract_pptx_theme(bad) is None


def test_pptx_write_applies_uploaded_template_theme(tmp_path, monkeypatch):
    from tools.document_tools import pptx_write

    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(tmp_path))
    tpl = tmp_path / "校规模板.pptx"
    _write_theme_pptx(tpl)
    out = tmp_path / "deck.pptx"
    captured: dict = {}
    result = json.loads(
        pptx_write(
            path=str(out),
            title="按校模板生成",
            slides=[{"title": "研究背景", "bullets": ["信息爆炸"]}],
            template="paper_report",
            visual_master="signal",
            template_path=str(tpl),
            callback=_fake_render_callback(captured),
        )
    )
    assert result["ok"] is True
    deck = captured["artifact"]["deck"]
    # The uploaded template's theme overrides the selected built-in master.
    assert deck["visual_master_palette"]["accent"] == "E6FF3D"
    assert deck["visual_master_palette"]["fonts"]["major"] == "Source Han Sans"
    assert deck["visual_master_name"] == "校规模板"
    # The built-in selection is still carried for fallback fields.
    assert deck["visual_master"] == "signal"


def test_normalize_deck_slides_keeps_single_richest_agenda():
    from tools.document_tools import _normalize_deck_slides

    slides = _normalize_deck_slides([
        {"slide_type": "agenda", "title": "封面信息", "bullets": ["作者：某某"]},
        {"slide_type": "agenda", "title": "汇报提纲", "bullets": ["背景", "方法", "总结"]},
        {"slide_type": "claim_bullets", "title": "研究背景", "bullets": ["信息爆炸"]},
    ])
    agendas = [s for s in slides if s["slide_type"] == "agenda"]
    assert len(agendas) == 1, "exactly one agenda must survive"
    # The richer agenda (more bullets) is the one kept.
    assert agendas[0]["title"] == "汇报提纲"
    assert [s["title"] for s in slides] == ["汇报提纲", "研究背景"]


def test_normalize_deck_slides_drops_empty_and_exact_duplicates():
    from tools.document_tools import _normalize_deck_slides

    slides = _normalize_deck_slides([
        {"slide_type": "claim_bullets", "title": "", "bullets": []},  # structurally empty -> drop
        {"slide_type": "claim_bullets", "title": "应用场景", "bullets": ["舆情监测", "品牌营销"]},
        {"slide_type": "claim_bullets", "title": "应用场景", "bullets": ["舆情监测", "品牌营销"]},  # exact dup -> drop
    ])
    assert [s["title"] for s in slides] == ["应用场景"]


def test_normalize_deck_slides_preserves_distinct_placeholders():
    from tools.document_tools import _normalize_deck_slides

    # Two screenshot placeholders with no bullets must both survive — a repeated
    # "insert result here" cue is legitimate and must not be deduped away.
    slides = _normalize_deck_slides([
        {"slide_type": "screenshot_placeholder", "title": "运行结果", "bullets": [],
         "placeholder": {"label": "运行结果", "caption": "请插入截图"}},
        {"slide_type": "screenshot_placeholder", "title": "测试结果", "bullets": [],
         "placeholder": {"label": "测试结果", "caption": "请插入截图"}},
    ])
    assert len(slides) == 2


def test_pptx_write_deck_spec_carries_structured_slide_types(tmp_path):
    from tools.document_tools import pptx_write

    out = tmp_path / "structured-student.pptx"
    captured: dict = {}
    result = json.loads(
        pptx_write(
            path=str(out),
            title="高质量可交付 PPT",
            slides=[
                {"slide_type": "agenda", "title": "汇报提纲", "bullets": ["研究背景", "系统设计"]},
                {
                    "slide_type": "diagram",
                    "title": "系统架构",
                    "diagram": {"nodes": ["Vue 前端", "Spring Boot API", "MySQL 数据库"]},
                },
                {
                    "slide_type": "table",
                    "title": "测试用例汇总",
                    "table": {
                        "headers": ["模块", "用例", "结果"],
                        "rows": [["仪表盘", "时间筛选", "通过"]],
                    },
                },
                {
                    "slide_type": "screenshot_placeholder",
                    "title": "系统运行截图",
                    "placeholder": {
                        "label": "待放入仪表盘截图",
                        "caption": "展示电、水、气总览",
                        "source_hint": "从系统首页截取真实运行画面。",
                    },
                    "notes": "备用讲稿。",
                },
            ],
            template="paper_report",
            callback=_fake_render_callback(captured),
        )
    )

    assert result["ok"] is True
    slides = captured["artifact"]["deck"]["slides"]
    assert slides[0]["slide_type"] == "agenda"
    assert slides[1]["diagram"]["nodes"] == ["Vue 前端", "Spring Boot API", "MySQL 数据库"]
    assert slides[2]["table"]["headers"] == ["模块", "用例", "结果"]
    assert slides[2]["table"]["rows"] == [["仪表盘", "时间筛选", "通过"]]
    assert slides[3]["placeholder"]["source_hint"] == "从系统首页截取真实运行画面。"
    assert slides[3]["notes"] == "备用讲稿。"


def test_pptx_write_passes_valid_layout_hint_and_drops_unknown(tmp_path):
    from tools.document_tools import pptx_write

    captured: dict = {}
    json.loads(
        pptx_write(
            path=str(tmp_path / "layout.pptx"),
            title="版式提示",
            slides=[
                {"title": "对比", "bullets": ["A", "B"], "layout": "comparison_cards"},
                {"title": "乱填", "bullets": ["x"], "layout": "not_a_layout"},
            ],
            template="course_report",
            callback=_fake_render_callback(captured),
        )
    )
    slides = captured["artifact"]["deck"]["slides"]
    assert slides[0]["layout"] == "comparison_cards"
    # Unknown layout hints are dropped so the renderer auto-selects.
    assert "layout" not in slides[1]


def test_pptx_write_unknown_slide_type_normalizes_to_claim_bullets(tmp_path):
    from tools.document_tools import pptx_write

    captured: dict = {}
    json.loads(
        pptx_write(
            path=str(tmp_path / "unknown.pptx"),
            title="未知类型",
            slides=[{"slide_type": "mystery", "title": "仍然生成", "bullets": ["回退"]}],
            template="course_report",
            callback=_fake_render_callback(captured),
        )
    )
    assert captured["artifact"]["deck"]["slides"][0]["slide_type"] == "claim_bullets"


def test_pptx_write_normalizes_unknown_template_and_visual_master(tmp_path):
    from tools.document_tools import pptx_write

    captured: dict = {}
    result = json.loads(
        pptx_write(
            path=str(tmp_path / "fallback.pptx"),
            title="测试",
            slides=[{"title": "页", "bullets": ["内容"]}],
            template="unknown_style",
            visual_master="unknown-master",
            callback=_fake_render_callback(captured),
        )
    )
    assert result["template"] == "course_report"
    assert result["visual_master"] == "default_native"
    assert result["visual_master_name"] == "Default native renderer"
    assert captured["artifact"]["deck"]["template"] == "course_report"
    assert captured["artifact"]["deck"]["visual_master"] == "default_native"


def test_pptx_write_requires_interactive_callback():
    from tools.document_tools import pptx_write

    result = json.loads(
        pptx_write(
            path="deck.pptx",
            title="无 UI",
            slides=[{"title": "页", "bullets": ["内容"]}],
            callback=None,
        )
    )
    assert result["error"]
    assert result["code"] == "pptx_render_unavailable"


def test_pptx_write_propagates_webview_render_error(tmp_path):
    from tools.document_tools import pptx_write

    def err_cb(question, choices, kind=None, artifact=None):
        return {"action": "error", "text": "PptxGenJS boom", "data": {}}

    result = json.loads(
        pptx_write(
            path=str(tmp_path / "deck.pptx"),
            title="出错",
            slides=[{"title": "页", "bullets": ["内容"]}],
            callback=err_cb,
        )
    )
    assert result["error"] == "PptxGenJS boom"
    assert result["code"] == "pptx_render_failed"


def test_pptx_write_rejects_output_outside_workspace(tmp_path, monkeypatch):
    from tools.document_tools import pptx_write

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(workspace))

    out = tmp_path / "outside.pptx"
    captured: dict = {}
    result = json.loads(
        pptx_write(
            path=str(out),
            title="越界",
            slides=[{"title": "页", "bullets": ["内容"]}],
            callback=_fake_render_callback(captured),
        )
    )
    assert result["code"] == "outside_workspace"
    assert not out.exists()


def test_pdf_write_generates_pdf_and_html_sidecar(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    rendered_specs = []

    def fake_render(spec):
        rendered_specs.append(spec)
        return b"%PDF-1.4\nfake pdf\n%%EOF", 2, "reportlab_pdf_v1"

    monkeypatch.setattr(document_tools, "_render_pdf_with_reportlab", fake_render)

    out = tmp_path / "course-report.pdf"
    result = json.loads(
        document_tools.pdf_write(
            path=str(out),
            title="课程总结报告",
            document={
                "sections": [
                    {
                        "title": "研究目标",
                        "paragraphs": ["把材料整理成可提交 PDF。"],
                        "bullets": ["保留表格", "保留公式"],
                    },
                    {
                        "title": "结果汇总",
                        "table": {
                            "headers": ["模块", "结果"],
                            "rows": [["PDF writer", "通过"]],
                        },
                    },
                ],
            },
            template="academic_report",
        )
    )

    assert result["ok"] is True
    assert result["path"] == str(out)
    assert result["html_path"] == str(out.with_suffix(".html"))
    assert result["page_count"] == 2
    assert result["renderer"] == "reportlab_pdf_v1"
    assert out.read_bytes().startswith(b"%PDF-1.4")
    html = out.with_suffix(".html").read_text(encoding="utf-8")
    assert "课程总结报告" in html
    assert "研究目标" in html
    assert "PDF writer" in html
    assert rendered_specs[0]["template"] == "academic_report"
    assert rendered_specs[0]["blocks"][0]["type"] == "heading"


def test_pdf_write_rejects_output_outside_workspace(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(workspace))
    monkeypatch.setattr(
        document_tools,
        "_render_pdf_with_reportlab",
        lambda _spec: (b"%PDF-1.4\nfake\n%%EOF", 1, "reportlab_pdf_v1"),
    )

    out = tmp_path / "outside.pdf"
    result = json.loads(
        document_tools.pdf_write(
            path=str(out),
            title="越界 PDF",
            document={"blocks": [{"type": "paragraph", "text": "内容"}]},
        )
    )

    assert result["code"] == "outside_workspace"
    assert not out.exists()
    assert not out.with_suffix(".html").exists()


def test_html_write_generates_standalone_document(tmp_path):
    import tools.document_tools as document_tools

    out = tmp_path / "web-report"  # suffix added by the writer
    result = json.loads(
        document_tools.html_write(
            path=str(out),
            title="网页学习报告",
            document={
                "sections": [
                    {
                        "title": "知识结构",
                        "paragraphs": ["把材料整理成可分享的网页。"],
                        "bullets": ["响应式布局", "自带样式"],
                    },
                    {
                        "title": "结果汇总",
                        "table": {"headers": ["模块", "结果"], "rows": [["HTML writer", "通过"]]},
                    },
                ],
            },
            template="code_report",
        )
    )

    written = tmp_path / "web-report.html"
    assert result["ok"] is True
    assert result["path"] == str(written)
    assert result["renderer"] == "standalone_html_v1"
    assert result["block_count"] > 0
    html = written.read_text(encoding="utf-8")
    # First-class HTML, not the print sidecar: responsive container + viewport.
    assert "<main class=\"container\">" in html
    assert "viewport" in html
    assert "网页学习报告" in html
    assert "知识结构" in html
    assert "HTML writer" in html


def test_html_write_coerces_json_string_document(tmp_path):
    """LLM tool-calls often stringify the document arg; it must still render DOM.

    Regression: a JSON-string ``document`` was dumped verbatim into one paragraph
    (raw JSON on the page), and a dict whose ``blocks`` was a JSON string fell
    through to an empty document (block_count 1, only the title).
    """
    import tools.document_tools as document_tools

    document = {
        "blocks": [
            {"type": "heading", "text": "结构标题"},
            {"type": "bullets", "items": ["第一点", "第二点"]},
            {"type": "table", "headers": ["列"], "rows": [["值"]]},
        ]
    }

    # Whole document passed as a JSON string.
    out = tmp_path / "stringified"
    result = json.loads(
        document_tools.html_write(
            path=str(out),
            title="字符串文档",
            document=json.dumps(document, ensure_ascii=False),
        )
    )
    html = (tmp_path / "stringified.html").read_text(encoding="utf-8")
    assert result["block_count"] == 3
    assert "<h2>结构标题</h2>" in html
    assert "<li>第一点</li>" in html
    assert "<table>" in html
    assert '"type": "heading"' not in html  # raw JSON must not leak into the page

    # Dict whose blocks value is itself a JSON string.
    out2 = tmp_path / "nested-string"
    result2 = json.loads(
        document_tools.html_write(
            path=str(out2),
            title="嵌套字符串",
            document={"blocks": json.dumps(document["blocks"], ensure_ascii=False)},
        )
    )
    assert result2["block_count"] == 3


def test_html_write_repairs_unescaped_quotes_in_json_string(tmp_path):
    """PDF-sourced content carries content quotes the model forgets to escape.

    Regression: the model stringifies ``document`` and embeds raw ``"`` inside a
    value (e.g. 信息"大爆炸"时代), so strict json.loads aborts and the whole
    malformed JSON was dumped onto the page as one paragraph. The writer must
    repair the unescaped quotes and still render real DOM.
    """
    import tools.document_tools as document_tools

    malformed = (
        '{"sections": [{"title": "研究背景", '
        '"paragraphs": ["信息时代进入"信息大爆炸"阶段，数据持续增长。"], '
        '"bullets": ["第一点", "第二点"]}]}'
    )
    # Sanity: the payload really is invalid JSON as written.
    with pytest.raises(json.JSONDecodeError):
        json.loads(malformed)

    out = tmp_path / "pdf-sourced"
    result = json.loads(
        document_tools.html_write(path=str(out), title="PDF 报告", document=malformed)
    )
    html = (tmp_path / "pdf-sourced.html").read_text(encoding="utf-8")
    assert result["block_count"] >= 3
    assert "<h2>研究背景</h2>" in html
    assert "<li>第一点</li>" in html
    assert "信息大爆炸" in html  # content quote text survives
    assert '"sections":' not in html  # raw JSON must not leak


def test_html_write_rejects_output_outside_workspace(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(workspace))

    out = tmp_path / "outside.html"
    result = json.loads(
        document_tools.html_write(
            path=str(out),
            title="越界 HTML",
            document={"blocks": [{"type": "paragraph", "text": "内容"}]},
        )
    )

    assert result["code"] == "outside_workspace"
    assert not out.exists()


def test_docx_write_generates_editable_word_document(tmp_path):
    import tools.document_tools as document_tools

    out = tmp_path / "word-report"  # suffix added by the writer
    result = json.loads(
        document_tools.docx_write(
            path=str(out),
            title="Word 学习报告",
            document={
                "sections": [
                    {
                        "title": "知识结构",
                        "paragraphs": ["把材料整理成可继续编辑的 Word。"],
                        "bullets": ["保留表格", "保留要点"],
                    },
                    {
                        "title": "结果汇总",
                        "table": {"headers": ["模块", "结果"], "rows": [["DOCX writer", "通过"]]},
                    },
                ],
            },
            template="code_report",
        )
    )

    written = tmp_path / "word-report.docx"
    assert result["ok"] is True
    assert result["path"] == str(written)
    assert result["renderer"] == "python_docx_v1"
    assert result["block_count"] > 0

    # Round-trip with python-docx to prove it is a real, editable .docx.
    from docx import Document

    doc = Document(str(written))
    texts = [p.text for p in doc.paragraphs]
    assert any(t == "Word 学习报告" for t in texts)
    assert any("保留表格" in t for t in texts)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2  # header + one data row


def test_docx_write_rejects_output_outside_workspace(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(workspace))

    out = tmp_path / "outside.docx"
    result = json.loads(
        document_tools.docx_write(
            path=str(out),
            title="越界 Word",
            document={"blocks": [{"type": "paragraph", "text": "内容"}]},
        )
    )

    assert result["code"] == "outside_workspace"
    assert not out.exists()


def test_pdf_fast_text_path_skips_docling_for_text_pdf(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    pdf = tmp_path / "thesis.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    monkeypatch.setattr(
        document_tools,
        "_read_pdf_with_pypdf",
        lambda p: {"ok": True, "engine": "pypdf", "mode": "fallback", "path": str(p), "pages": 3, "content": "正文内容 " * 400},
    )

    def _no_docling(_p, _mode):
        raise AssertionError("Docling must not run for a text-rich PDF in auto mode")

    monkeypatch.setattr(document_tools, "_read_with_docling", _no_docling)

    result = json.loads(document_tools.document_read_precise(path=str(pdf), mode="auto"))

    assert result["ok"] is True
    assert result["engine"] == "pypdf"
    assert "Fast text-only PDF read" in result["warning"]


def test_pdf_fast_text_path_falls_through_to_docling_for_scanned_pdf(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    monkeypatch.setattr(
        document_tools,
        "_read_pdf_with_pypdf",
        lambda p: {"ok": True, "engine": "pypdf", "pages": 5, "content": "   "},
    )
    monkeypatch.setattr(
        document_tools,
        "_read_with_docling",
        lambda p, mode: {"ok": True, "engine": "docling", "mode": mode, "path": str(p), "pages": 5, "content": "Docling layout text"},
    )

    result = json.loads(document_tools.document_read_precise(path=str(pdf), mode="auto"))

    assert result["ok"] is True
    assert result["engine"] == "docling"


def test_pdf_precise_mode_still_uses_docling(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    pdf = tmp_path / "precise.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")

    def _no_pypdf(_p):
        raise AssertionError("precise mode should not take the fast pypdf path")

    monkeypatch.setattr(document_tools, "_read_pdf_with_pypdf", _no_pypdf)
    monkeypatch.setattr(
        document_tools,
        "_read_with_docling",
        lambda p, mode: {"ok": True, "engine": "docling", "mode": mode, "path": str(p), "pages": 2, "content": "precise"},
    )

    result = json.loads(document_tools.document_read_precise(path=str(pdf), mode="precise"))

    assert result["ok"] is True
    assert result["engine"] == "docling"


def test_pdf_read_precise_falls_back_to_pypdf(tmp_path):
    from pypdf import PdfWriter
    from tools.document_tools import pdf_read_precise

    pdf = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with pdf.open("wb") as f:
        writer.write(f)

    result = json.loads(pdf_read_precise(path=str(pdf)))

    assert result["ok"] is True
    assert result["engine"] in {"docling", "pypdf"}
    assert result["pages"] == 1
    if result["engine"] == "pypdf":
        assert "Docling" in result["warning"]
        assert "docling_error" in result


def test_docling_converter_is_cached(monkeypatch):
    import tools.document_tools as document_tools

    calls = []

    class FakeConverter:
        pass

    def fake_create(profile="fast"):
        calls.append("create")
        return FakeConverter()

    document_tools.reset_docling_converter_cache()
    monkeypatch.setattr(document_tools, "_create_docling_converter", fake_create)

    first = document_tools._get_docling_converter()
    second = document_tools._get_docling_converter()

    assert first is second
    assert calls == ["create"]
    document_tools.reset_docling_converter_cache()


def test_docling_profile_for_mode_defaults_to_fast():
    from tools.document_tools import _docling_profile_for_mode

    assert _docling_profile_for_mode("auto") == "fast"
    assert _docling_profile_for_mode("") == "fast"
    assert _docling_profile_for_mode("precise") == "precise"


def test_docling_profile_for_mode_supports_math():
    from tools.document_tools import _docling_profile_for_mode

    assert _docling_profile_for_mode("math") == "math"
    assert _docling_profile_for_mode("MATH") == "math"


def test_configure_pdf_pipeline_options_enables_formula_for_math():
    import tools.document_tools as document_tools

    class Options:
        do_ocr = True
        do_table_structure = True
        do_code_enrichment = False
        do_formula_enrichment = False

    fast = Options()
    document_tools._configure_pdf_pipeline_options(fast, "fast")
    assert fast.do_formula_enrichment is False
    assert fast.do_code_enrichment is False

    math = Options()
    document_tools._configure_pdf_pipeline_options(math, "math")
    assert math.do_formula_enrichment is True
    assert math.do_code_enrichment is False


def test_require_math_artifacts_raises_without_codeformula(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    bundle = tmp_path / "bundle"
    (bundle / "docling-models").mkdir(parents=True)
    monkeypatch.setenv("HERMESDESK_BUNDLE_DIR", str(bundle))
    monkeypatch.delenv("DOCLING_ARTIFACTS_PATH", raising=False)

    with pytest.raises(ValueError, match="mode=math requires offline CodeFormula"):
        document_tools._require_math_artifacts_bundled()


def test_ensure_math_artifacts_uses_desktop_first_use_download(monkeypatch):
    import docling_math_models
    import tools.document_tools as document_tools

    calls = []

    def fake_ensure():
        calls.append("ensure")

    monkeypatch.setattr(docling_math_models, "ensure_code_formula_available_for_math", fake_ensure)

    document_tools._ensure_math_artifacts()

    assert calls == ["ensure"]


def test_format_docling_error_surfaces_settings_hint_for_missing_model():
    from tools.document_tools import _format_docling_error

    try:
        from docling_math_models import CodeFormulaMissingError
    except ImportError:
        pytest.skip("docling_math_models not on path")

    msg = _format_docling_error(
        CodeFormulaMissingError(
            "code_formula_model_missing: mode=math requires ds4sd/CodeFormula (~500 MB). "
            "Download in Settings."
        )
    )
    assert "code_formula_model_missing" in msg
    assert "Settings" in msg


def test_read_document_precise_math_mode_does_not_fallback_on_missing_model(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    source = tmp_path / "formula.pdf"
    source.write_bytes(b"%PDF-1.4 fake")
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(tmp_path))

    try:
        from docling_math_models import CodeFormulaMissingError
    except ImportError:
        pytest.skip("docling_math_models not on path")

    def fake_read(_path: Path, _mode: str):
        raise CodeFormulaMissingError(
            "code_formula_model_missing: mode=math requires ds4sd/CodeFormula (~500 MB). "
            "Download in Settings."
        )

    monkeypatch.setattr(document_tools, "_read_with_docling", fake_read)

    result = json.loads(document_tools.document_read_precise(path=str(source), mode="math"))

    assert result["ok"] is False
    assert result["code"] == "docling_math_unavailable"
    assert "code_formula_model_missing" in result["docling_error"]
    assert "Settings" in result["docling_error"]


def test_prime_torch_keeps_existing_modules_on_failure(monkeypatch):
    import sys
    import types
    import tools.document_tools as document_tools

    fake_torch = types.ModuleType("torch")
    fake_child = types.ModuleType("torch.partial")
    monkeypatch.setitem(sys.modules, "torch", fake_torch)
    monkeypatch.setitem(sys.modules, "torch.partial", fake_child)

    real_import = __import__

    def fake_import(name, *args, **kwargs):
        if name == "torch":
            return fake_torch
        if name == "torch.library":
            raise AttributeError("partially initialized module 'torch' has no attribute 'library'")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", fake_import)

    with pytest.raises(AttributeError):
        document_tools._prime_torch_for_docling()

    assert sys.modules["torch"] is fake_torch
    assert sys.modules["torch.partial"] is fake_child


def test_document_read_precise_routes_docx_through_docling(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    docx = tmp_path / "brief.docx"
    docx.write_bytes(b"fake docx")

    def fake_read(path: Path, mode: str):
        return {
            "ok": True,
            "engine": "docling",
            "mode": mode,
            "path": str(path),
            "pages": 0,
            "content": "Docling DOCX content",
        }

    monkeypatch.setattr(document_tools, "_read_with_docling", fake_read)

    result = json.loads(document_tools.document_read_precise(path=str(docx), mode="precise"))

    assert result["ok"] is True
    assert result["engine"] == "docling"
    assert result["mode"] == "precise"
    assert result["content"] == "Docling DOCX content"


def test_document_read_precise_math_mode_uses_docling_even_for_lightweight_suffix(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    source = tmp_path / "formula.md"
    source.write_text("Euler: $e^{i\\pi}+1=0$", encoding="utf-8")
    bundle = tmp_path / "bundle"
    formula = bundle / "docling-models" / "ds4sd--CodeFormula"
    formula.mkdir(parents=True)
    (formula / "model.safetensors").write_bytes(b"x")
    monkeypatch.setenv("HERMESDESK_BUNDLE_DIR", str(bundle))
    monkeypatch.setenv("HERMESDESK_WORKSPACE", str(tmp_path))
    monkeypatch.delenv("DOCLING_ARTIFACTS_PATH", raising=False)

    def fake_read(path: Path, mode: str):
        return {
            "ok": True,
            "engine": "docling",
            "mode": mode,
            "profile": "math",
            "path": str(path),
            "content": "$$e^{i\\pi}+1=0$$",
        }

    monkeypatch.setattr(document_tools, "_read_with_docling", fake_read)

    result = json.loads(document_tools.document_read_precise(path=str(source), mode="math"))

    assert result["ok"] is True
    assert result["engine"] == "docling"
    assert result["profile"] == "math"
    assert result["metadata"]["kind"] == "markdown"
    assert "$$e^{i\\pi}+1=0$$" in result["content"]


def test_document_read_precise_rejects_legacy_doc_with_hint(tmp_path):
    from tools.document_tools import document_read_precise

    doc = tmp_path / "legacy.doc"
    doc.write_bytes(b"not really a Word binary")

    result = json.loads(document_read_precise(path=str(doc)))

    assert result["ok"] is False
    assert result["code"] == "unsupported_legacy_doc"
    assert "LibreOffice" in result["hint"]


def test_run_on_docling_thread_serializes_calls():
    from tools import document_tools

    seen: list[int] = []

    def _record(value: int) -> int:
        seen.append(value)
        return value

    assert document_tools._run_on_docling_thread(_record, 1) == 1
    assert document_tools._run_on_docling_thread(_record, 2) == 2
    assert seen == [1, 2]


def test_format_docling_error_surfaces_network_failures():
    from tools.document_tools import _format_docling_error

    msg = _format_docling_error(
        ConnectionError("Connection to huggingface.co timed out.")
    )
    assert "Docling model load failed" in msg
    assert "huggingface.co" in msg


def test_format_docling_error_flags_unsupported_torch_runtime():
    from tools.document_tools import _format_docling_error

    msg = _format_docling_error(
        AttributeError(
            "partially initialized module 'torch' has no attribute 'library' "
            "(most likely due to a circular import)"
        )
    )
    assert "environment problem" in msg
    assert "3.11" in msg


def test_format_docling_error_flags_pdfium_page_count_failure():
    from tools.document_tools import _format_docling_error

    msg = _format_docling_error(ValueError("Inconsistent number of pages: 73!=-1"))
    assert "environment problem" in msg
    assert "3.11" in msg


def test_format_docling_error_flags_duplicate_torch_kernel_registration():
    from tools.document_tools import _format_docling_error

    msg = _format_docling_error(
        RuntimeError(
            "This is not allowed since there's already a kernel registered from python "
            "overriding wait_tensor's behavior for Autograd dispatch key and "
            "_c10d_functional namespace."
        )
    )
    assert "environment problem" in msg
    assert "restart" in msg.lower()


def test_format_docling_error_surfaces_path_policy_blocks():
    from tools.document_tools import _format_docling_error

    msg = _format_docling_error(
        PermissionError(
            "Kabuqina path policy blocked read to \\\\.\\nul\\ "
            "(allowed root: C:\\\\Users\\\\X13\\\\Documents\\\\KabuqinaWork)"
        )
    )
    assert "path policy" in msg.lower() or "PathPolicy" in msg or "blocked read" in msg


def test_resolve_docling_artifacts_path_from_bundle_env(tmp_path, monkeypatch):
    from tools.document_tools import _resolve_docling_artifacts_path

    models = tmp_path / "docling-models"
    models.mkdir()
    monkeypatch.setenv("HERMESDESK_BUNDLE_DIR", str(tmp_path))
    monkeypatch.delenv("DOCLING_ARTIFACTS_PATH", raising=False)

    assert _resolve_docling_artifacts_path() == models


def test_document_read_precise_fast_text_skips_docling_and_writes_read_cache(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    data_dir = tmp_path / "data"
    source = tmp_path / "notes.md"
    source.write_text("# 标题\n\nRead 层应该快速读取 markdown。", encoding="utf-8")

    monkeypatch.setenv("HERMESDESK_DATA_DIR", str(data_dir))

    def fail_docling(path: Path, mode: str):
        raise AssertionError("fast markdown reads should not initialize Docling")

    monkeypatch.setattr(document_tools, "_read_with_docling", fail_docling)

    result = json.loads(document_tools.document_read_precise(path=str(source), mode="auto"))

    assert result["ok"] is True
    assert result["engine"] == "text"
    assert result["read_id"]
    assert result["cache_path"].startswith(str(data_dir))
    assert result["metadata"]["kind"] == "markdown"
    assert result["content"].startswith("# 标题")


def test_document_read_precise_include_content_false_returns_cache_handle(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    monkeypatch.setenv("HERMESDESK_DATA_DIR", str(tmp_path / "data"))
    source = tmp_path / "brief.txt"
    source.write_text("这是一份较长材料的正文。", encoding="utf-8")

    result = json.loads(
        document_tools.document_read_precise(path=str(source), mode="fast", include_content=False)
    )

    assert result["ok"] is True
    assert result["read_id"]
    assert result["content"] == ""
    assert result["content_omitted"] is True
    assert Path(result["cache_path"]).exists()


def test_document_read_precise_large_output_surfaces_cache_hint_before_content(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    monkeypatch.setenv("HERMESDESK_DATA_DIR", str(tmp_path / "data"))
    source = tmp_path / "formula.pdf"
    source.write_bytes(b"%PDF-1.4\n")
    payload = {
        "ok": True,
        "engine": "docling",
        "mode": "math",
        "content": "x" * 120_000,
    }

    result = document_tools._finalize_read_payload(payload, source, include_content=True)
    serialized = document_tools._json(result)

    assert result["content_hint"]
    assert "read_file" in result["content_hint"]
    assert "vision_analyze" in result["content_hint"]
    assert serialized.index('"content_hint"') < serialized.index('"content"')


def test_pdf_read_precise_uses_common_read_pipeline_for_pdf(tmp_path, monkeypatch):
    import tools.document_tools as document_tools

    pdf = tmp_path / "paper.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")

    calls = []

    def fake_read(path: Path, *, mode: str, include_content: bool):
        calls.append((path, mode, include_content))
        return {"ok": True, "path": str(path), "engine": "fixture", "content": "PDF text"}

    monkeypatch.setattr(document_tools, "_read_document_precise_payload", fake_read)

    result = json.loads(document_tools.pdf_read_precise(path=str(pdf), mode="precise", include_content=False))

    assert result["ok"] is True
    assert result["content"] == "PDF text"
    assert calls == [(pdf, "precise", False)]


def test_read_tool_schemas_advertise_math_mode():
    from tools.document_tools import DOCUMENT_READ_PRECISE_SCHEMA, PDF_READ_PRECISE_SCHEMA, PDF_WRITE_SCHEMA

    assert "math" in PDF_READ_PRECISE_SCHEMA["parameters"]["properties"]["mode"]["description"]
    assert "math" in DOCUMENT_READ_PRECISE_SCHEMA["parameters"]["properties"]["mode"]["description"]
    assert PDF_WRITE_SCHEMA["name"] == "pdf_write"
    assert "HTML" in PDF_WRITE_SCHEMA["description"]
