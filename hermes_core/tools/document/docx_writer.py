# Copyright 2026 Kabuqina Contributors
# SPDX-License-Identifier: Apache-2.0

"""DOCX document writer."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any, Dict

from tools.registry import tool_error
from tools.document.common import (
    DocumentSpecError,
    _json,
    _text,
    _validate_write_path,
)
from tools.document.spec import _build_pdf_spec, _document_spec_error

logger = logging.getLogger(__name__)


def _docx_add_block(doc: Any, block: Dict[str, Any]) -> None:
    """Render one normalized block into a python-docx Document."""
    from docx.shared import Pt

    block_type = block.get("type")
    if block_type == "heading":
        level = max(1, min(4, int(block.get("level") or 2)))
        doc.add_heading(_text(block.get("text")), level=level)
    elif block_type == "bullets":
        for item in block.get("items") or []:
            doc.add_paragraph(_text(item), style="List Bullet")
    elif block_type == "table":
        headers = [_text(h) for h in block.get("headers") or []]
        rows = [[_text(c) for c in row] for row in block.get("rows") or []]
        cols = len(headers) or (len(rows[0]) if rows else 0)
        if cols:
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            if headers:
                cells = table.add_row().cells
                for idx, header in enumerate(headers[:cols]):
                    cells[idx].text = header
                    for paragraph in cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[:cols]):
                    cells[idx].text = value
    elif block_type == "code":
        language = _text(block.get("language"))
        if language:
            caption = doc.add_paragraph()
            caption_run = caption.add_run(language)
            caption_run.italic = True
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(_text(block.get("text")))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    elif block_type == "formula":
        paragraph = doc.add_paragraph()
        paragraph.add_run(_text(block.get("text"))).italic = True
    elif block_type == "image_placeholder":
        label = _text(block.get("label")) or "Image placeholder"
        caption = _text(block.get("caption"))
        source_hint = _text(block.get("source_hint"))
        note = f"[{label}]"
        if caption:
            note += f" {caption}"
        if source_hint:
            note += f"（替换为：{source_hint}）"
        paragraph = doc.add_paragraph()
        paragraph.add_run(note).italic = True
    elif block_type == "page_break":
        doc.add_page_break()
    else:
        doc.add_paragraph(_text(block.get("text")))


def _render_docx(spec: Dict[str, Any]) -> bytes:
    """Render the normalized document spec into .docx bytes via python-docx.

    Consumes the same structured blocks as the PDF and HTML writers, so a single
    reviewed outline can target .docx, .pdf, or .html.
    """
    from docx import Document

    doc = Document()
    title = _text(spec.get("title"))
    if title:
        doc.add_heading(title, level=0)
    subtitle = " · ".join(
        part for part in (_text(spec.get("template_name")), _text(spec.get("template_subtitle"))) if part
    )
    if subtitle:
        paragraph = doc.add_paragraph()
        paragraph.add_run(subtitle).italic = True
    for block in spec.get("blocks") or []:
        _docx_add_block(doc, block)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def docx_write(
    path: str,
    title: str,
    document: Any,
    template: str = "academic_report",
    visual_master: str = "default_docx",
) -> str:
    """Create an editable .docx deliverable from a structured document spec.

    Consumes the same structured document/blocks contract as pdf_write and
    html_write (sections or blocks of heading, paragraph, bullets, table, code,
    formula, image_placeholder, page_break), so the reader -> material_index ->
    planner layers can target Word the same way they target PDF/HTML.
    """
    if not _text(path):
        return tool_error("docx_write requires an output path.", code="missing_path")
    out = Path(path).expanduser()
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")

    validation_error = _validate_write_path(out, path)
    if validation_error is not None:
        return validation_error

    try:
        spec = _build_pdf_spec(title, document, template, visual_master)
    except DocumentSpecError as exc:
        return _document_spec_error(exc)
    try:
        docx_bytes = _render_docx(spec)
    except ImportError as exc:
        return tool_error(
            f"DOCX rendering requires python-docx in the desktop Python bundle: {exc}",
            code="docx_render_unavailable",
        )
    except Exception as exc:
        return tool_error(f"DOCX rendering failed: {exc}", code="docx_render_failed")

    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)
    except OSError as exc:
        return tool_error(f"DOCX write failed: {exc}", code="docx_write_failed")

    return _json(
        {
            "ok": True,
            "path": str(out),
            "template": spec["template"],
            "visual_master": spec["visual_master"],
            "renderer": "python_docx_v1",
            "block_count": len(spec.get("blocks") or []),
            "bytes": len(docx_bytes),
        }
    )
