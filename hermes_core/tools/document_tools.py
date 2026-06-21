"""Document read/write tools for student deliverables."""

from __future__ import annotations

import json
import logging
import os
import sys
import hashlib
import html
import io
import tempfile
import threading
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from tools.deliverable_contract import slide_layout_set, slide_type_set
from tools.registry import registry, tool_error
from tools.document.common import (
    _text, _list, _string_list, _dict, DocumentSpecError, Rgb,
    _PDF_BLOCK_TYPES, _PDF_TEMPLATES, _PPTX_SLIDE_TYPES, _PPTX_SLIDE_LAYOUTS,
    _desktop_workspace_root, _is_outside_workspace, _json,
    _power_user_reads_anywhere, _validate_read_path,
)
from tools.document.reading import (
    warm_docling_converter, reset_docling_converter_cache,
    pdf_read_precise, document_read_precise,
)
from tools.document.latex_render import (
    _LATEX_SYMBOLS, _latex_group, _latex_atom,
    _latex_command_html, _render_latex_html, _formula_to_html,
)

logger = logging.getLogger(__name__)









# Slide-object vocabulary is owned by tools/deliverable_contract.py — the single
# source of truth shared with the planner prompt (agent/prompt_builder.py) so the
# writer normalizes against the exact set the planner is told to emit.

# Optional per-slide layout hints the planner may set; the web renderer
# (renderDeck.ts) auto-selects by content when omitted. Keep in sync with
# SLIDE_LAYOUT_IDS in web/src/chat/pptx/renderDeck.ts.

















































































































@dataclass(frozen=True)
class _PptxTheme:
    key: str
    subtitle: str
    badge: str
    bg: Rgb
    title: Rgb
    subtitle_color: Rgb
    body: Rgb
    accent: Rgb
    accent_light: Rgb
    title_band: Optional[Rgb] = None
    footer_band: Optional[Rgb] = None


_PPTX_THEMES: Dict[str, _PptxTheme] = {
    "course_report": _PptxTheme(
        key="course_report",
        subtitle="课程汇报",
        badge="课程报告",
        bg=(239, 246, 255),
        title=(30, 64, 175),
        subtitle_color=(37, 99, 235),
        body=(51, 65, 85),
        accent=(37, 99, 235),
        accent_light=(147, 197, 253),
        title_band=(37, 99, 235),
    ),
    "paper_report": _PptxTheme(
        key="paper_report",
        subtitle="论文 / 文献汇报",
        badge="论文报告",
        bg=(255, 251, 235),
        title=(20, 83, 45),
        subtitle_color=(22, 101, 52),
        body=(68, 64, 60),
        accent=(22, 101, 52),
        accent_light=(134, 239, 172),
    ),
    "code_defense": _PptxTheme(
        key="code_defense",
        subtitle="课程设计答辩",
        badge="项目答辩",
        bg=(241, 245, 249),
        title=(49, 46, 129),
        subtitle_color=(234, 88, 12),
        body=(71, 85, 105),
        accent=(99, 102, 241),
        accent_light=(249, 115, 22),
        footer_band=(30, 41, 59),
    ),
}

_PPTX_VISUAL_MASTERS: Dict[str, Dict[str, str]] = {
    "default_native": {"name": "Default native renderer", "dir": ""},
    "soft_editorial": {"name": "Soft Editorial", "dir": "soft-editorial"},
    "blue_professional": {"name": "Blue Professional", "dir": "blue-professional"},
    "signal": {"name": "Signal", "dir": "signal"},
    "neo_grid_bold": {"name": "Neo Grid Bold", "dir": "neo-grid-bold"},
    "editorial_forest": {"name": "Editorial Forest", "dir": "editorial-forest"},
}


def _normalize_pptx_template(template: str) -> str:
    key = (template or "course_report").strip().lower()
    return key if key in _PPTX_THEMES else "course_report"


def _normalize_pptx_visual_master(visual_master: str) -> str:
    key = (visual_master or "default_native").strip().lower().replace("-", "_")
    return key if key in _PPTX_VISUAL_MASTERS else "default_native"


def _pptx_visual_master_name(visual_master: str) -> str:
    return _PPTX_VISUAL_MASTERS[visual_master]["name"]


def _get_pptx_theme(template: str) -> _PptxTheme:
    return _PPTX_THEMES[_normalize_pptx_template(template)]


def _normalize_slide_type(raw: Any) -> str:
    key = str(raw or "claim_bullets").strip().lower()
    return key if key in _PPTX_SLIDE_TYPES else "claim_bullets"










def _validate_write_path(out_path: Path, original_path: str) -> Optional[str]:
    """Workspace guard for *write* targets (the file need not exist yet)."""
    workspace = _desktop_workspace_root()
    if workspace is None:
        return None
    try:
        resolved = out_path.resolve()
    except OSError:
        resolved = out_path
    if _is_outside_workspace(resolved, workspace):
        return tool_error(
            f"Output path is outside the Kabuqina workspace ({workspace}): {original_path}",
            code="outside_workspace",
            workspace=str(workspace),
            hint="Write the output file into the workspace (or a subfolder of it).",
        )
    return None


def _deck_slide_spec(raw: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize one planner slide into a renderer-friendly entry.

    The web layer (PptxGenJS) renders from this spec, so keep it structured and
    bounded — same field contract the old python-pptx renderer consumed.
    """
    raw = _dict(raw)
    slide_type = _normalize_slide_type(raw.get("slide_type"))
    entry: Dict[str, Any] = {
        "slide_type": slide_type,
        "title": _text(raw.get("title")),
        "subtitle": _text(raw.get("subtitle")),
        "bullets": _string_list(raw.get("bullets"), limit=8),
        "notes": _text(raw.get("notes")),
        "tags": _string_list(raw.get("tags"), limit=6),
    }
    # Optional per-slide layout hint; the web renderer auto-selects by content
    # when this is absent or unknown (see renderDeck.ts:chooseLayout).
    layout = _text(raw.get("layout"))
    if layout in _PPTX_SLIDE_LAYOUTS:
        entry["layout"] = layout
    diagram = _dict(raw.get("diagram"))
    if diagram:
        entry["diagram"] = {
            "nodes": _string_list(diagram.get("nodes") or diagram.get("steps"), limit=8),
        }
    table = _dict(raw.get("table"))
    if table:
        headers = _string_list(table.get("headers"), limit=6)
        entry["table"] = {
            "headers": headers,
            "rows": [
                _string_list(row, limit=len(headers) or 6)
                for row in _list(table.get("rows"))[:8]
            ],
        }
    placeholder = _dict(raw.get("placeholder"))
    if placeholder:
        entry["placeholder"] = {
            "label": _text(placeholder.get("label")),
            "caption": _text(placeholder.get("caption")),
            "source_hint": _text(placeholder.get("source_hint")),
        }
    return entry


def _slide_has_body(slide: Dict[str, Any]) -> bool:
    """A slide carries content if it has bullets, a table, a diagram, or a placeholder."""
    return bool(
        slide.get("bullets")
        or slide.get("table")
        or slide.get("diagram")
        or slide.get("placeholder")
    )


def _normalize_deck_slides(slides: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Deterministic structural cleanup applied after per-slide normalization.

    Model-independent guardrails that guarantee a sound deck skeleton regardless
    of planner quality. We only *remove* clearly-broken structure — never rewrite
    or invent content, which stays the planner/model's job. This is the hard
    backstop paired with the soft guidance in
    ``prompt_builder.build_deliverable_planner_prompt`` ("Slide content quality"):
    the prompt aims the model high, this catches the misses for free.

    Three high-confidence rules, each only ever a deletion:

    1. Drop structurally-empty slides (no title and no body of any kind).
    2. Exactly one agenda — keep the richest (most bullets; ties → earliest) and
       drop the rest. Two agenda slides is never intended.
    3. Drop exact-duplicate *content* slides (same type + title + bullets),
       keeping the first. Placeholder slides are never deduped, since a repeated
       "insert screenshot here" cue can be legitimate in a defense deck.
    """
    cleaned = [s for s in slides if s.get("title") or _slide_has_body(s)]

    agenda_idx = [i for i, s in enumerate(cleaned) if s.get("slide_type") == "agenda"]
    if len(agenda_idx) > 1:
        keep = max(agenda_idx, key=lambda i: (len(cleaned[i].get("bullets") or []), -i))
        cleaned = [
            s for i, s in enumerate(cleaned)
            if s.get("slide_type") != "agenda" or i == keep
        ]

    seen: set = set()
    deduped: List[Dict[str, Any]] = []
    for s in cleaned:
        signature = (s.get("slide_type"), s.get("title"), tuple(s.get("bullets") or []))
        if s.get("bullets") and signature in seen:
            continue
        seen.add(signature)
        deduped.append(s)
    return deduped


def _deck_meta(raw: Any) -> Dict[str, str]:
    """Deck-level presentation metadata that belongs on the cover, not in slides.

    Giving author / affiliation / date / citation a structural home removes the
    reason a planner would otherwise cram author info into an agenda bullet.
    """
    raw = _dict(raw)
    meta: Dict[str, str] = {}
    for key in ("author", "affiliation", "date", "citation"):
        value = _text(raw.get(key))
        if value:
            meta[key] = value
    return meta


def _hex6(value: str) -> str:
    """Normalize a DrawingML colour to a 6-hex string, or '' if not parseable."""
    v = (value or "").strip().lstrip("#").upper()
    if len(v) == 6 and all(c in "0123456789ABCDEF" for c in v):
        return v
    return ""


def _theme_color(scheme_el: Any, tag: str, ns: Dict[str, str]) -> str:
    """Read one <a:clrScheme> entry (srgbClr val=... or sysClr lastClr=...)."""
    node = scheme_el.find(f"a:{tag}", ns)
    if node is None:
        return ""
    srgb = node.find("a:srgbClr", ns)
    if srgb is not None:
        return _hex6(srgb.get("val", ""))
    sysclr = node.find("a:sysClr", ns)
    if sysclr is not None:
        return _hex6(sysclr.get("lastClr", ""))
    return ""


def _extract_pptx_theme(template_path: Path) -> Optional[Dict[str, Any]]:
    """Derive an inline visual master (palette + fonts) from an uploaded .pptx.

    Route A of template support: rather than rendering *into* the school template
    (which PptxGenJS cannot do), we read its DrawingML theme and reuse its colours
    and fonts so the generated deck matches the school's look while keeping our
    own rich layouts. Pure stdlib (zip + XML), deterministic, no dependency on the
    python-pptx writer. Returns ``None`` if the theme cannot be parsed, so the
    caller falls back to the selected built-in master.
    """
    import xml.etree.ElementTree as ET
    import zipfile

    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    try:
        with zipfile.ZipFile(template_path) as zf:
            with zf.open("ppt/theme/theme1.xml") as fh:
                root = ET.parse(fh).getroot()
    except (KeyError, OSError, zipfile.BadZipFile, ET.ParseError):
        return None

    elements = root.find("a:themeElements", ns)
    if elements is None:
        return None
    scheme = elements.find("a:clrScheme", ns)
    fonts_el = elements.find("a:fontScheme", ns)
    if scheme is None:
        return None

    # lt1/dk1 are usually window bg/text; dk2/accent1 carry the brand identity.
    background = _theme_color(scheme, "lt1", ns) or _theme_color(scheme, "lt2", ns)
    title = _theme_color(scheme, "dk2", ns) or _theme_color(scheme, "dk1", ns)
    body = _theme_color(scheme, "dk1", ns) or title
    accent = _theme_color(scheme, "accent1", ns)
    accent2 = _theme_color(scheme, "accent2", ns) or accent

    palette: Dict[str, Any] = {}
    if background:
        palette["background"] = background
    if title:
        palette["title"] = title
    if body and body != background:
        palette["body"] = body
    if accent:
        palette["accent"] = accent
    if accent2:
        palette["accent2"] = accent2

    if fonts_el is not None:
        def _typeface(role: str) -> str:
            font = fonts_el.find(f"a:{role}", ns)
            latin = font.find("a:latin", ns) if font is not None else None
            face = (latin.get("typeface", "") if latin is not None else "").strip()
            # "+mj-lt"/"+mn-lt" are theme self-references, not real font names.
            return face if face and not face.startswith("+") else ""
        major = _typeface("majorFont")
        minor = _typeface("minorFont")
        if major or minor:
            palette["fonts"] = {"major": major or minor, "minor": minor or major}

    # Require at least an accent or a usable background to consider it a theme.
    if not (palette.get("accent") or palette.get("background")):
        return None
    return palette


def _build_deck_spec(
    title: str,
    slides: List[Dict[str, Any]],
    theme: "_PptxTheme",
    visual_master: str,
    meta: Any = None,
    theme_override: Optional[Dict[str, Any]] = None,
    template_name: str = "",
) -> Dict[str, Any]:
    normalized = _normalize_deck_slides([_deck_slide_spec(raw) for raw in (slides or [])])
    spec: Dict[str, Any] = {
        "title": _text(title) or "学生汇报",
        "template": theme.key,
        "template_subtitle": theme.subtitle,
        "template_badge": theme.badge,
        "visual_master": visual_master,
        "visual_master_name": _pptx_visual_master_name(visual_master),
        "page_size": {"width": 13.333, "height": 7.5},
        "meta": _deck_meta(meta),
        "slides": normalized,
    }
    if theme_override:
        spec["visual_master_palette"] = theme_override
        if template_name:
            spec["visual_master_name"] = template_name
    return spec


def _normalize_pdf_template(template: str) -> str:
    key = _text(template).lower()
    return key if key in _PDF_TEMPLATES else "academic_report"


def _pdf_block(raw: Any) -> Dict[str, Any]:
    if isinstance(raw, str):
        return {"type": "paragraph", "text": raw}
    raw = _dict(raw)
    block_type = _text(raw.get("type") or raw.get("block_type")).lower()
    if block_type not in _PDF_BLOCK_TYPES:
        block_type = "paragraph"
    if block_type == "heading":
        level = raw.get("level", 2)
        try:
            level = max(1, min(4, int(level)))
        except (TypeError, ValueError):
            level = 2
        return {"type": "heading", "level": level, "text": _text(raw.get("text") or raw.get("title"))}
    if block_type == "bullets":
        return {"type": "bullets", "items": _string_list(raw.get("items") or raw.get("bullets"), limit=24)}
    if block_type == "table":
        table = _dict(raw.get("table")) or raw
        headers = _string_list(table.get("headers"), limit=8)
        rows = [_string_list(row, limit=len(headers) or 8) for row in _list(table.get("rows"))[:40]]
        return {"type": "table", "headers": headers, "rows": rows}
    if block_type == "code":
        return {
            "type": "code",
            "language": _text(raw.get("language")),
            "text": _text(raw.get("text") or raw.get("code")),
        }
    if block_type == "formula":
        return {"type": "formula", "text": _text(raw.get("text") or raw.get("latex") or raw.get("formula"))}
    if block_type == "image_placeholder":
        return {
            "type": "image_placeholder",
            "label": _text(raw.get("label") or raw.get("title") or "Image placeholder"),
            "caption": _text(raw.get("caption")),
            "source_hint": _text(raw.get("source_hint")),
        }
    if block_type == "page_break":
        return {"type": "page_break"}
    return {"type": "paragraph", "text": _text(raw.get("text") or raw.get("content") or raw.get("body"))}


def _pdf_section_blocks(section: Dict[str, Any]) -> List[Dict[str, Any]]:
    section = _dict(section)
    blocks: List[Dict[str, Any]] = []
    title = _text(section.get("title") or section.get("heading"))
    if title:
        blocks.append({"type": "heading", "level": 2, "text": title})
    for paragraph in _list(section.get("paragraphs")):
        if _text(paragraph):
            blocks.append({"type": "paragraph", "text": _text(paragraph)})
    content = section.get("content")
    if isinstance(content, list):
        blocks.extend(_pdf_block(item) for item in content)
    elif _text(content):
        blocks.append({"type": "paragraph", "text": _text(content)})
    bullets = _string_list(section.get("bullets"), limit=24)
    if bullets:
        blocks.append({"type": "bullets", "items": bullets})
    if _dict(section.get("table")):
        blocks.append(_pdf_block({"type": "table", "table": section.get("table")}))
    if _text(section.get("code")):
        blocks.append(_pdf_block({"type": "code", "code": section.get("code"), "language": section.get("language")}))
    if _text(section.get("formula") or section.get("latex")):
        blocks.append(_pdf_block({"type": "formula", "text": section.get("formula") or section.get("latex")}))
    placeholder = _dict(section.get("image_placeholder") or section.get("placeholder"))
    if placeholder:
        blocks.append(_pdf_block({"type": "image_placeholder", **placeholder}))
    return blocks


def _repair_jsonish(text: str) -> str:
    """Best-effort repair of the JSON mistakes LLMs make when stringifying args.

    Single left-to-right pass that tracks whether we are inside a string literal:

    * Unescaped ``"`` inside a value — the dominant failure mode. Source text
      lifted from a PDF often contains content quotes (e.g. 信息"大爆炸"时代);
      when the model hand-builds the ``document`` JSON it forgets to escape them,
      so strict ``json.loads`` aborts at the first one. A ``"`` is treated as the
      string terminator only when the next significant char is structural
      (``:`` ``,`` ``}`` ``]`` or end-of-input); otherwise it is content and gets
      escaped to ``\\"``.
    * Trailing commas before ``}`` / ``]`` are dropped.

    Escape sequences are passed through verbatim so already-valid input is left
    unchanged. This is a heuristic, not a parser: the caller still runs
    ``json.loads`` on the result and falls back to the raw string if it fails.
    """
    out: List[str] = []
    i = 0
    n = len(text)
    in_str = False
    while i < n:
        c = text[i]
        if not in_str:
            if c == ",":
                j = i + 1
                while j < n and text[j] in " \t\r\n":
                    j += 1
                if j < n and text[j] in "}]":
                    i += 1  # drop trailing comma
                    continue
            out.append(c)
            if c == '"':
                in_str = True
            i += 1
            continue
        if c == "\\":
            out.append(text[i:i + 2])
            i += 2
            continue
        if c == '"':
            j = i + 1
            while j < n and text[j] in " \t\r\n":
                j += 1
            nxt = text[j] if j < n else ""
            if nxt in (":", ",", "}", "]", ""):
                out.append('"')
                in_str = False
            else:
                out.append('\\"')
            i += 1
            continue
        out.append(c)
        i += 1
    return "".join(out)


def _coerce_json_container(value: Any) -> Any:
    """Decode a JSON-string into the dict/list it encodes; pass anything else through.

    LLM tool-calls routinely stringify nested object/array arguments, so the
    ``document`` (or its ``blocks`` / ``sections``) can arrive as a JSON string
    even though the schema declares an object. Without this, ``_pdf_blocks``
    treated the whole JSON string as a single paragraph (dumping raw JSON onto
    the page) or fell through to an empty document. Genuine prose stays a string:
    only text that starts with ``[``/``{`` is treated as a container candidate.

    When strict parsing fails (commonly unescaped content quotes from PDF text),
    one repair pass via ``_repair_jsonish`` is attempted before giving up.
    """
    if isinstance(value, str):
        text = value.strip()
        if text[:1] in ("[", "{"):
            for candidate in (text, _repair_jsonish(text)):
                try:
                    parsed = json.loads(candidate)
                except (ValueError, TypeError):
                    continue
                if isinstance(parsed, (dict, list)):
                    return parsed
            raise DocumentSpecError(
                "document looks like JSON but could not be parsed. Pass document as a real "
                "object, or fix the malformed JSON string before calling the writer."
            )
    return value


def _document_spec_error(exc: DocumentSpecError) -> str:
    return tool_error(str(exc), ok=False, code="invalid_document_json")


def _pdf_blocks(document: Any) -> List[Dict[str, Any]]:
    document = _coerce_json_container(document)
    if isinstance(document, str):
        return [{"type": "paragraph", "text": document}]
    if isinstance(document, list):
        return [_pdf_block(item) for item in document]
    document = _dict(document)
    blocks_value = _coerce_json_container(document.get("blocks"))
    if isinstance(blocks_value, list):
        return [_pdf_block(item) for item in blocks_value]
    blocks: List[Dict[str, Any]] = []
    for section in _list(_coerce_json_container(document.get("sections"))):
        blocks.extend(_pdf_section_blocks(_dict(section)))
    if blocks:
        return blocks
    for key in ("summary", "abstract", "content", "body", "text"):
        value = document.get(key)
        if _text(value):
            blocks.append({"type": "paragraph", "text": _text(value)})
    return blocks or [{"type": "paragraph", "text": ""}]


def _build_pdf_spec(
    title: str,
    document: Any,
    template: str,
    visual_master: str,
) -> Dict[str, Any]:
    template_key = _normalize_pdf_template(template)
    return {
        "title": _text(title) or "Kabuqina PDF",
        "template": template_key,
        "template_name": _PDF_TEMPLATES[template_key]["title"],
        "template_subtitle": _PDF_TEMPLATES[template_key]["subtitle"],
        "visual_master": _text(visual_master) or "default_print",
        "page_size": "A4",
        "blocks": _pdf_blocks(document),
    }














def _block_to_html(block: Dict[str, Any]) -> str:
    block_type = block.get("type")
    if block_type == "heading":
        level = max(1, min(4, int(block.get("level") or 2)))
        return f"<h{level}>{html.escape(_text(block.get('text')))}</h{level}>"
    if block_type == "bullets":
        items = "".join(f"<li>{html.escape(item)}</li>" for item in block.get("items") or [])
        return f"<ul>{items}</ul>"
    if block_type == "table":
        headers = "".join(f"<th>{html.escape(item)}</th>" for item in block.get("headers") or [])
        rows = []
        for row in block.get("rows") or []:
            rows.append("<tr>" + "".join(f"<td>{html.escape(item)}</td>" for item in row) + "</tr>")
        head = f"<thead><tr>{headers}</tr></thead>" if headers else ""
        return f"<table>{head}<tbody>{''.join(rows)}</tbody></table>"
    if block_type == "code":
        language = html.escape(_text(block.get("language")))
        label = f"<figcaption>{language}</figcaption>" if language else ""
        return f"<figure class=\"code\">{label}<pre>{html.escape(_text(block.get('text')))}</pre></figure>"
    if block_type == "formula":
        return f"<div class=\"formula\">{_formula_to_html(_text(block.get('text')))}</div>"
    if block_type == "image_placeholder":
        label = html.escape(_text(block.get("label")))
        caption = html.escape(_text(block.get("caption")))
        source_hint = html.escape(_text(block.get("source_hint")))
        return (
            "<figure class=\"image-placeholder\">"
            f"<div>{label}</div><figcaption>{caption}</figcaption><small>{source_hint}</small>"
            "</figure>"
        )
    if block_type == "page_break":
        return "<div class=\"page-break\"></div>"
    return f"<p>{html.escape(_text(block.get('text')))}</p>"


def _build_pdf_html(spec: Dict[str, Any]) -> str:
    accent = _PDF_TEMPLATES.get(spec["template"], _PDF_TEMPLATES["academic_report"])["accent"]
    accent_css = f"rgb({accent[0]}, {accent[1]}, {accent[2]})"
    body = "\n".join(_block_to_html(block) for block in spec.get("blocks") or [])
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{html.escape(_text(spec.get("title")))}</title>
  <style>
    @page {{ size: A4; margin: 18mm; }}
    body {{ color: #111827; font-family: "Microsoft YaHei", "Noto Sans CJK SC", Arial, sans-serif; line-height: 1.55; }}
    h1 {{ color: {accent_css}; font-size: 26px; margin: 0 0 6px; }}
    .subtitle {{ color: #64748b; margin: 0 0 22px; }}
    h2, h3, h4 {{ color: #111827; margin: 22px 0 8px; }}
    p {{ margin: 8px 0; }}
    ul {{ margin: 8px 0 12px 22px; }}
    table {{ border-collapse: collapse; margin: 12px 0; width: 100%; }}
    th, td {{ border: 1px solid #cbd5e1; padding: 7px 9px; text-align: left; vertical-align: top; }}
    th {{ background: #f1f5f9; }}
    pre {{ background: #0f172a; color: #e2e8f0; border-radius: 6px; padding: 12px; white-space: pre-wrap; }}
    figure {{ margin: 12px 0; }}
    .formula {{ background: #f8fafc; border-left: 4px solid {accent_css}; padding: 10px 12px; font-family: Cambria Math, serif; }}
    .formula-math {{ font-size: 1.05em; word-spacing: 0.08em; }}
    .frac {{ display: inline-flex; flex-direction: column; align-items: center; vertical-align: middle; line-height: 1.05; margin: 0 0.12em; }}
    .frac .num {{ border-bottom: 1px solid currentColor; padding: 0 0.18em 1px; }}
    .frac .den {{ padding: 1px 0.18em 0; }}
    .sqrt::before {{ content: "√"; font-size: 1.12em; }}
    .sqrt .radicand {{ border-top: 1px solid currentColor; padding: 0 0.12em; }}
    .overline {{ text-decoration: overline; text-decoration-thickness: 1px; }}
    .vec-mark {{ display: inline-block; margin-left: -0.1em; transform: translateY(-0.18em); }}
    .image-placeholder {{ border: 1px dashed #94a3b8; border-radius: 6px; padding: 18px; color: #475569; }}
    .page-break {{ break-after: page; page-break-after: always; }}
  </style>
</head>
<body>
  <h1>{html.escape(_text(spec.get("title")))}</h1>
  <p class="subtitle">{html.escape(_text(spec.get("template_name")))} · {html.escape(_text(spec.get("template_subtitle")))}</p>
  {body}
</body>
</html>
"""


def _count_pdf_pages(pdf_bytes: bytes) -> int:
    try:
        from pypdf import PdfReader

        return len(PdfReader(io.BytesIO(pdf_bytes)).pages)
    except Exception:
        return 0


def _render_pdf_from_html(html_source: str) -> Tuple[bytes, int, str]:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise ImportError(
            "HTML print PDF rendering requires the Python Playwright package."
        ) from exc

    launch_attempts: List[str] = []
    with sync_playwright() as pw:
        launch_options: List[Dict[str, Any]] = []
        browser_path = _text(os.getenv("HERMESDESK_PDF_BROWSER_PATH"))
        if browser_path:
            launch_options.append({"executable_path": browser_path, "headless": True})

        channel = _text(os.getenv("HERMESDESK_PDF_BROWSER_CHANNEL"))
        if channel:
            launch_options.append({"channel": channel, "headless": True})
        elif sys.platform.startswith("win"):
            launch_options.append({"channel": "msedge", "headless": True})

        launch_options.append({"headless": True})

        browser = None
        for options in launch_options:
            try:
                browser = pw.chromium.launch(**options)
                break
            except Exception as exc:
                launch_attempts.append(f"{options}: {exc}")
        if browser is None:
            detail = "; ".join(launch_attempts) or "no launch attempts"
            raise RuntimeError(f"Could not launch Chromium for PDF printing: {detail}")

        try:
            page = browser.new_page(viewport={"width": 794, "height": 1123})
            page.set_content(html_source, wait_until="load")
            page.emulate_media(media="print")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                prefer_css_page_size=True,
            )
        finally:
            browser.close()

    return pdf_bytes, _count_pdf_pages(pdf_bytes), "chromium_print_v1"


def render_pdf_from_html_source(html_source: str) -> Tuple[bytes, int, str]:
    """Render an arbitrary trusted HTML print source using the core PDF renderer."""
    return _render_pdf_from_html(html_source)


def _build_standalone_html(spec: Dict[str, Any]) -> str:
    """Screen-first standalone HTML deliverable.

    Reuses the same normalized blocks and per-block rendering as the PDF sidecar
    (``_block_to_html``), but wraps them in a responsive, centered container with
    both screen and print styling so the .html is a first-class output, not just
    a print source.
    """
    accent = _PDF_TEMPLATES.get(spec["template"], _PDF_TEMPLATES["academic_report"])["accent"]
    accent_css = f"rgb({accent[0]}, {accent[1]}, {accent[2]})"
    body = "\n".join(_block_to_html(block) for block in spec.get("blocks") or [])
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(_text(spec.get("title")))}</title>
  <style>
    :root {{ --accent: {accent_css}; }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: #f1f5f9; color: #111827;
      font-family: "Microsoft YaHei", "Noto Sans CJK SC", Arial, sans-serif; line-height: 1.6; }}
    main.container {{ max-width: 820px; margin: 0 auto; padding: 40px 28px 64px;
      background: #ffffff; min-height: 100vh; box-shadow: 0 1px 24px rgba(15, 23, 42, 0.08); }}
    h1 {{ color: var(--accent); font-size: 28px; margin: 0 0 6px; }}
    .subtitle {{ color: #64748b; margin: 0 0 26px; border-bottom: 1px solid #e2e8f0; padding-bottom: 16px; }}
    h2, h3, h4 {{ color: #0f172a; margin: 26px 0 8px; }}
    p {{ margin: 10px 0; }}
    ul {{ margin: 10px 0 14px 22px; }}
    table {{ border-collapse: collapse; margin: 14px 0; width: 100%; }}
    th, td {{ border: 1px solid #cbd5e1; padding: 8px 10px; text-align: left; vertical-align: top; }}
    th {{ background: #f1f5f9; }}
    figure {{ margin: 14px 0; }}
    figure.code figcaption {{ color: #64748b; font-size: 13px; margin-bottom: 4px; }}
    pre {{ background: #0f172a; color: #e2e8f0; border-radius: 8px; padding: 14px; overflow-x: auto; white-space: pre-wrap; }}
    .formula {{ background: #f8fafc; border-left: 4px solid var(--accent); padding: 12px 14px; font-family: Cambria Math, serif; }}
    .formula-math {{ font-size: 1.05em; word-spacing: 0.08em; }}
    .frac {{ display: inline-flex; flex-direction: column; align-items: center; vertical-align: middle; line-height: 1.05; margin: 0 0.12em; }}
    .frac .num {{ border-bottom: 1px solid currentColor; padding: 0 0.18em 1px; }}
    .frac .den {{ padding: 1px 0.18em 0; }}
    .sqrt::before {{ content: "√"; font-size: 1.12em; }}
    .sqrt .radicand {{ border-top: 1px solid currentColor; padding: 0 0.12em; }}
    .overline {{ text-decoration: overline; text-decoration-thickness: 1px; }}
    .vec-mark {{ display: inline-block; margin-left: -0.1em; transform: translateY(-0.18em); }}
    .image-placeholder {{ border: 1px dashed #94a3b8; border-radius: 8px; padding: 20px; color: #475569; text-align: center; }}
    .page-break {{ break-after: page; page-break-after: always; }}
    @media (max-width: 640px) {{ main.container {{ padding: 24px 16px 48px; }} }}
    @media print {{ body {{ background: #fff; }} main.container {{ box-shadow: none; max-width: none; }} }}
  </style>
</head>
<body>
  <main class="container">
    <h1>{html.escape(_text(spec.get("title")))}</h1>
    <p class="subtitle">{html.escape(_text(spec.get("template_name")))} · {html.escape(_text(spec.get("template_subtitle")))}</p>
    {body}
  </main>
</body>
</html>
"""


def _docx_add_block(doc: Any, block: Dict[str, Any]) -> None:
    """Render one normalized block into a python-docx Document."""
    from docx.shared import Pt

    block_type = block.get("type")
    if block_type == "heading":
        level = max(1, min(4, int(block.get("level") or 2)))
        doc.add_heading(_text(block.get("text")), level=level)
    elif block_type == "bullets":
        for item in block.get("items") or []:
            doc.add_paragraph(_text(item), style="List Bullet")
    elif block_type == "table":
        headers = [_text(h) for h in block.get("headers") or []]
        rows = [[_text(c) for c in row] for row in block.get("rows") or []]
        cols = len(headers) or (len(rows[0]) if rows else 0)
        if cols:
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            if headers:
                cells = table.add_row().cells
                for idx, header in enumerate(headers[:cols]):
                    cells[idx].text = header
                    for paragraph in cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[:cols]):
                    cells[idx].text = value
    elif block_type == "code":
        language = _text(block.get("language"))
        if language:
            caption = doc.add_paragraph()
            caption_run = caption.add_run(language)
            caption_run.italic = True
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(_text(block.get("text")))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    elif block_type == "formula":
        paragraph = doc.add_paragraph()
        paragraph.add_run(_text(block.get("text"))).italic = True
    elif block_type == "image_placeholder":
        label = _text(block.get("label")) or "Image placeholder"
        caption = _text(block.get("caption"))
        source_hint = _text(block.get("source_hint"))
        note = f"[{label}]"
        if caption:
            note += f" {caption}"
        if source_hint:
            note += f"（替换为：{source_hint}）"
        paragraph = doc.add_paragraph()
        paragraph.add_run(note).italic = True
    elif block_type == "page_break":
        doc.add_page_break()
    else:
        doc.add_paragraph(_text(block.get("text")))


def _render_docx(spec: Dict[str, Any]) -> bytes:
    """Render the normalized document spec into .docx bytes via python-docx.

    Consumes the same structured blocks as the PDF and HTML writers, so a single
    reviewed outline can target .docx, .pdf, or .html.
    """
    from docx import Document

    doc = Document()
    title = _text(spec.get("title"))
    if title:
        doc.add_heading(title, level=0)
    subtitle = " · ".join(
        part for part in (_text(spec.get("template_name")), _text(spec.get("template_subtitle"))) if part
    )
    if subtitle:
        paragraph = doc.add_paragraph()
        paragraph.add_run(subtitle).italic = True
    for block in spec.get("blocks") or []:
        _docx_add_block(doc, block)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _wrap_pdf_text(text: str, *, width: int = 78) -> List[str]:
    lines: List[str] = []
    for raw_line in _text(text).splitlines() or [""]:
        if not raw_line:
            lines.append("")
            continue
        lines.extend(textwrap.wrap(raw_line, width=width, break_long_words=True, replace_whitespace=False) or [""])
    return lines


def _render_pdf_with_reportlab(spec: Dict[str, Any]) -> Tuple[bytes, int, str]:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"

    template = _PDF_TEMPLATES.get(spec["template"], _PDF_TEMPLATES["academic_report"])
    accent = template["accent"]
    buffer = io.BytesIO()
    page_width, page_height = A4
    margin = 1.8 * cm
    line_height = 14
    page_count = 1
    y = page_height - margin
    c = canvas.Canvas(buffer, pagesize=A4)

    def new_page() -> None:
        nonlocal y, page_count
        c.showPage()
        page_count += 1
        y = page_height - margin

    def ensure(space: float = line_height) -> None:
        if y - space < margin:
            new_page()

    def draw_text(text: str, *, size: int = 10, leading: int = line_height, indent: float = 0, font: str = font_name) -> None:
        nonlocal y
        for line in _wrap_pdf_text(text, width=88 if size <= 10 else 64):
            ensure(leading)
            c.setFont(font, size)
            c.drawString(margin + indent, y, line)
            y -= leading

    c.setTitle(_text(spec.get("title")))
    c.setAuthor("Kabuqina")
    c.setFillColorRGB(accent[0] / 255, accent[1] / 255, accent[2] / 255)
    draw_text(_text(spec.get("title")), size=18, leading=24)
    c.setFillColorRGB(0.39, 0.45, 0.55)
    draw_text(f"{spec.get('template_name')} · {spec.get('template_subtitle')}", size=9, leading=16)
    c.setStrokeColorRGB(accent[0] / 255, accent[1] / 255, accent[2] / 255)
    c.line(margin, y, page_width - margin, y)
    y -= 18
    c.setFillColorRGB(0.07, 0.09, 0.15)

    for block in spec.get("blocks") or []:
        block_type = block.get("type")
        if block_type == "page_break":
            new_page()
            continue
        if block_type == "heading":
            level = int(block.get("level") or 2)
            y -= 4
            c.setFillColorRGB(0.07, 0.09, 0.15)
            draw_text(_text(block.get("text")), size=15 if level <= 2 else 12, leading=20 if level <= 2 else 16)
            continue
        if block_type == "bullets":
            for item in block.get("items") or []:
                draw_text(f"• {item}", indent=10)
            y -= 4
            continue
        if block_type == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            if headers:
                draw_text(" | ".join(headers), size=9, leading=13)
            for row in rows:
                draw_text(" | ".join(row), size=9, leading=13)
            y -= 6
            continue
        if block_type == "code":
            language = _text(block.get("language"))
            if language:
                draw_text(language, size=8, leading=12)
            for line in _wrap_pdf_text(_text(block.get("text")), width=90):
                draw_text(line, size=8, leading=11, font="Courier")
            y -= 6
            continue
        if block_type == "formula":
            draw_text(_text(block.get("text")), size=11, leading=15)
            y -= 6
            continue
        if block_type == "image_placeholder":
            ensure(54)
            c.setStrokeColorRGB(0.58, 0.64, 0.72)
            c.rect(margin, y - 46, page_width - 2 * margin, 42, stroke=1, fill=0)
            c.setFillColorRGB(0.29, 0.33, 0.41)
            c.setFont(font_name, 9)
            c.drawString(margin + 10, y - 20, _text(block.get("label")))
            c.drawString(margin + 10, y - 34, _text(block.get("caption") or block.get("source_hint")))
            y -= 58
            c.setFillColorRGB(0.07, 0.09, 0.15)
            continue
        draw_text(_text(block.get("text")), size=10, leading=14)
        y -= 4

    c.save()
    return buffer.getvalue(), page_count, "reportlab_pdf_v1"


def pdf_write(
    path: str,
    title: str,
    document: Any,
    template: str = "academic_report",
    visual_master: str = "default_print",
) -> str:
    """Create a PDF and adjacent HTML source from a structured document spec."""
    if not _text(path):
        return tool_error("pdf_write requires an output path.", code="missing_path")
    out = Path(path).expanduser()
    if out.suffix.lower() != ".pdf":
        out = out.with_suffix(".pdf")

    validation_error = _validate_write_path(out, path)
    if validation_error is not None:
        return validation_error

    try:
        spec = _build_pdf_spec(title, document, template, visual_master)
    except DocumentSpecError as exc:
        return _document_spec_error(exc)
    html_source = _build_pdf_html(spec)
    warnings: List[str] = []
    try:
        pdf_bytes, page_count, renderer = _render_pdf_from_html(html_source)
    except Exception as html_exc:
        warnings.append(f"HTML print renderer unavailable; fell back to ReportLab: {html_exc}")
        try:
            pdf_bytes, page_count, renderer = _render_pdf_with_reportlab(spec)
        except ImportError as exc:
            return tool_error(
                (
                    "PDF rendering requires either Chromium/Playwright or ReportLab "
                    f"in the desktop Python bundle. HTML print error: {html_exc}; "
                    f"ReportLab error: {exc}"
                ),
                code="pdf_render_unavailable",
            )
        except Exception as exc:
            return tool_error(
                f"PDF rendering failed. HTML print error: {html_exc}; ReportLab error: {exc}",
                code="pdf_render_failed",
            )

    out.parent.mkdir(parents=True, exist_ok=True)
    html_path = out.with_suffix(".html")
    out.write_bytes(pdf_bytes)
    html_path.write_text(html_source, encoding="utf-8")
    payload = {
        "ok": True,
        "path": str(out),
        "html_path": str(html_path),
        "page_count": int(page_count),
        "template": spec["template"],
        "visual_master": spec["visual_master"],
        "renderer": renderer,
        "bytes": len(pdf_bytes),
    }
    if warnings:
        payload["warnings"] = warnings
    return _json(payload)


def html_write(
    path: str,
    title: str,
    document: Any,
    template: str = "academic_report",
    visual_master: str = "default_web",
) -> str:
    """Create a standalone .html deliverable from a structured document spec.

    This is the first-class HTML writer (not the print sidecar that pdf_write
    emits). It consumes the same structured document/blocks contract as pdf_write
    so the reader -> material_index -> planner layers can target either format,
    then renders a responsive, self-contained HTML page.
    """
    if not _text(path):
        return tool_error("html_write requires an output path.", code="missing_path")
    out = Path(path).expanduser()
    if out.suffix.lower() not in (".html", ".htm"):
        out = out.with_suffix(".html")

    validation_error = _validate_write_path(out, path)
    if validation_error is not None:
        return validation_error

    try:
        spec = _build_pdf_spec(title, document, template, visual_master)
    except DocumentSpecError as exc:
        return _document_spec_error(exc)
    html_source = _build_standalone_html(spec)

    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(html_source, encoding="utf-8")
    except OSError as exc:
        return tool_error(f"HTML write failed: {exc}", code="html_write_failed")

    return _json(
        {
            "ok": True,
            "path": str(out),
            "template": spec["template"],
            "visual_master": spec["visual_master"],
            "renderer": "standalone_html_v1",
            "block_count": len(spec.get("blocks") or []),
            "bytes": len(html_source.encode("utf-8")),
        }
    )


def docx_write(
    path: str,
    title: str,
    document: Any,
    template: str = "academic_report",
    visual_master: str = "default_docx",
) -> str:
    """Create an editable .docx deliverable from a structured document spec.

    Consumes the same structured document/blocks contract as pdf_write and
    html_write (sections or blocks of heading, paragraph, bullets, table, code,
    formula, image_placeholder, page_break), so the reader -> material_index ->
    planner layers can target Word the same way they target PDF/HTML.
    """
    if not _text(path):
        return tool_error("docx_write requires an output path.", code="missing_path")
    out = Path(path).expanduser()
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")

    validation_error = _validate_write_path(out, path)
    if validation_error is not None:
        return validation_error

    try:
        spec = _build_pdf_spec(title, document, template, visual_master)
    except DocumentSpecError as exc:
        return _document_spec_error(exc)
    try:
        docx_bytes = _render_docx(spec)
    except ImportError as exc:
        return tool_error(
            f"DOCX rendering requires python-docx in the desktop Python bundle: {exc}",
            code="docx_render_unavailable",
        )
    except Exception as exc:
        return tool_error(f"DOCX rendering failed: {exc}", code="docx_render_failed")

    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)
    except OSError as exc:
        return tool_error(f"DOCX write failed: {exc}", code="docx_write_failed")

    return _json(
        {
            "ok": True,
            "path": str(out),
            "template": spec["template"],
            "visual_master": spec["visual_master"],
            "renderer": "python_docx_v1",
            "block_count": len(spec.get("blocks") or []),
            "bytes": len(docx_bytes),
        }
    )


def pptx_write(
    path: str,
    title: str,
    slides: List[Dict[str, Any]],
    template: str = "course_report",
    visual_master: str = "default_native",
    meta: Optional[Dict[str, Any]] = None,
    template_path: str = "",
    callback=None,
) -> str:
    """Render a student deck with PptxGenJS in the desktop web layer.

    The agent loop (run_agent._invoke_tool) routes this tool with
    ``callback=self.clarify_callback``. We emit the deck spec as a
    ``kind="pptx_render"`` interaction; the webview builds the .pptx with
    PptxGenJS and returns it base64-encoded in the interaction ``data``, which
    we decode and write to the (workspace-validated) output path. There is no
    python-pptx writer anymore — Python only persists the bytes.
    """
    logger.info(
        "pptx_write called: callback_present=%s slides=%d template=%s visual_master=%s",
        callback is not None,
        len(slides or []),
        template,
        visual_master,
    )
    theme = _get_pptx_theme(template)
    selected_visual_master = _normalize_pptx_visual_master(visual_master)
    out = Path(path).expanduser()
    if out.suffix.lower() != ".pptx":
        out = out.with_suffix(".pptx")

    # Route A: an uploaded school template contributes its colours/fonts as an
    # inline visual master. Bad/unreadable templates degrade silently to the
    # selected built-in master rather than failing the whole deck.
    theme_override: Optional[Dict[str, Any]] = None
    template_name = ""
    if _text(template_path):
        tpl = Path(template_path).expanduser()
        read_error = _validate_read_path(tpl, template_path, "PPT template")
        if read_error is not None:
            return read_error
        theme_override = _extract_pptx_theme(tpl)
        if theme_override is not None:
            template_name = tpl.stem
            logger.info("pptx_write: applied uploaded template theme from %s", tpl.name)
        else:
            logger.warning("pptx_write: could not parse theme from %s; using built-in master", tpl.name)

    deck_spec = _build_deck_spec(
        title, slides, theme, selected_visual_master, meta, theme_override, template_name
    )

    if callback is None:
        return tool_error(
            "PptxGenJS rendering requires the Kabuqina desktop UI; no interactive "
            "renderer is available in this context.",
            code="pptx_render_unavailable",
        )

    artifact = {"type": "pptx_deck", "filename": out.name, "deck": deck_spec}
    question = f"生成 PPT：{deck_spec['title']}（{len(deck_spec['slides'])} 页内容）"
    try:
        response = callback(question, [], kind="pptx_render", artifact=artifact)
    except TypeError:
        return tool_error(
            "The active interaction callback does not support PptxGenJS rendering.",
            code="pptx_render_unsupported",
        )

    data: Dict[str, Any] = {}
    if isinstance(response, dict):
        action = str(response.get("action") or "")
        data = response.get("data") if isinstance(response.get("data"), dict) else {}
        if action in {"cancel", "timeout"}:
            return tool_error(
                f"PPT rendering was not completed (action={action}).",
                code="pptx_render_cancelled",
            )
        if action == "error":
            return tool_error(
                str(response.get("text") or "PptxGenJS rendering failed in the web layer."),
                code="pptx_render_failed",
            )

    b64 = str(data.get("pptx_base64") or "")
    if not b64:
        return tool_error(
            "The web renderer did not return a .pptx payload.",
            code="pptx_render_empty",
        )

    import base64

    try:
        raw_bytes = base64.b64decode(b64)
    except Exception as exc:
        return tool_error(f"Could not decode rendered .pptx: {exc}", code="pptx_render_decode_error")

    validation_error = _validate_write_path(out, path)
    if validation_error is not None:
        return validation_error
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(raw_bytes)

    slide_count = int(data.get("slide_count") or (len(deck_spec["slides"]) + 1))
    return _json(
        {
            "ok": True,
            "path": str(out),
            "slide_count": slide_count,
            "template": theme.key,
            "theme": theme.badge,
            "visual_master": selected_visual_master,
            "visual_master_name": _pptx_visual_master_name(selected_visual_master),
            "visual_master_renderer": "pptxgenjs_v1",
            "bytes": len(raw_bytes),
        }
    )


from tools.document.schemas import (
    PDF_READ_PRECISE_SCHEMA,
    DOCUMENT_READ_PRECISE_SCHEMA,
    PDF_WRITE_SCHEMA,
    HTML_WRITE_SCHEMA,
    DOCX_WRITE_SCHEMA,
    PPTX_WRITE_SCHEMA,
)
registry.register(
    name="pdf_read_precise",
    toolset="documents",
    schema=PDF_READ_PRECISE_SCHEMA,
    handler=lambda args, **kw: pdf_read_precise(
        path=args.get("path", ""),
        mode=args.get("mode", "auto"),
        include_content=bool(args.get("include_content", True)),
        page_start=args.get("page_start"),
        page_end=args.get("page_end"),
    ),
    check_fn=lambda: True,
    emoji="📄",
)

registry.register(
    name="document_read_precise",
    toolset="documents",
    schema=DOCUMENT_READ_PRECISE_SCHEMA,
    handler=lambda args, **kw: document_read_precise(
        path=args.get("path", ""),
        mode=args.get("mode", "auto"),
        include_content=bool(args.get("include_content", True)),
        page_start=args.get("page_start"),
        page_end=args.get("page_end"),
    ),
    check_fn=lambda: True,
    emoji="📚",
)

registry.register(
    name="pdf_write",
    toolset="documents",
    schema=PDF_WRITE_SCHEMA,
    handler=lambda args, **kw: pdf_write(
        path=args.get("path", ""),
        title=args.get("title", ""),
        document=args.get("document") or {},
        template=args.get("template", "academic_report"),
        visual_master=args.get("visual_master", "default_print"),
    ),
    check_fn=lambda: True,
    emoji="📝",
)

registry.register(
    name="html_write",
    toolset="documents",
    schema=HTML_WRITE_SCHEMA,
    handler=lambda args, **kw: html_write(
        path=args.get("path", ""),
        title=args.get("title", ""),
        document=args.get("document") or {},
        template=args.get("template", "academic_report"),
        visual_master=args.get("visual_master", "default_web"),
    ),
    check_fn=lambda: True,
    emoji="🌐",
)

registry.register(
    name="docx_write",
    toolset="documents",
    schema=DOCX_WRITE_SCHEMA,
    handler=lambda args, **kw: docx_write(
        path=args.get("path", ""),
        title=args.get("title", ""),
        document=args.get("document") or {},
        template=args.get("template", "academic_report"),
        visual_master=args.get("visual_master", "default_docx"),
    ),
    check_fn=lambda: True,
    emoji="📄",
)

registry.register(
    name="pptx_write",
    toolset="documents",
    schema=PPTX_WRITE_SCHEMA,
    handler=lambda args, **kw: pptx_write(
        path=args.get("path", ""),
        title=args.get("title", ""),
        slides=args.get("slides") or [],
        template=args.get("template", "course_report"),
        visual_master=args.get("visual_master", "default_native"),
        meta=args.get("meta"),
        template_path=args.get("template_path", ""),
        callback=kw.get("callback"),
    ),
    check_fn=lambda: True,
    emoji="📊",
)
